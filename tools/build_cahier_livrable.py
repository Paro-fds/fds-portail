from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, List, Tuple

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "cahier_des_charges.md"
OUT_DIR = ROOT / "livrables"
FIG_DIR = OUT_DIR / "figures"
DOCX_OUT = OUT_DIR / "Cahier_des_charges_FDS_Portail_Module_2.docx"
PDF_OUT = OUT_DIR / "Cahier_des_charges_FDS_Portail_Module_2.pdf"
LOGO = ROOT / "public" / "logo.jpg"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111C2D"
MUTED = "515F74"
GRID = "C8D0DC"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
GREEN = "1E7A5A"
RED = "B42318"
GOLD = "8A6A00"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> List[str]:
    words = text.split()
    lines: List[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded_box(
    draw: ImageDraw.ImageDraw,
    xy: Tuple[int, int, int, int],
    text: str,
    fill: str = "FFFFFF",
    outline: str = BLUE,
    color: str = INK,
    fnt: ImageFont.FreeTypeFont | None = None,
    radius: int = 14,
):
    fnt = fnt or font(24)
    draw.rounded_rectangle(xy, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=3)
    x1, y1, x2, y2 = xy
    lines = wrap(draw, text, fnt, x2 - x1 - 24)
    line_h = int(fnt.size * 1.25)
    total_h = len(lines) * line_h
    y = y1 + ((y2 - y1) - total_h) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2, y), line, fill=f"#{color}", font=fnt)
        y += line_h


def arrow(draw: ImageDraw.ImageDraw, start: Tuple[int, int], end: Tuple[int, int], color: str = MUTED, width: int = 3):
    draw.line([start, end], fill=f"#{color}", width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - direction * 14, ey - 8), (ex - direction * 14, ey + 8)]
    else:
        direction = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 8, ey - direction * 14), (ex + 8, ey - direction * 14)]
    draw.polygon(pts, fill=f"#{color}")


def label(draw: ImageDraw.ImageDraw, xy: Tuple[int, int], text: str, size: int = 22, color: str = INK, bold: bool = False):
    draw.text(xy, text, fill=f"#{color}", font=font(size, bold=bold))


def canvas(title: str, w: int = 1800, h: int = 1050) -> Tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, w, 76), fill=f"#{PALE_BLUE}")
    label(draw, (40, 20), title, 30, DARK_BLUE, True)
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    path = FIG_DIR / name
    img.save(path, "PNG")
    return path


def diagram_use_cases() -> Path:
    img, d = canvas("Figure 1 - Diagramme de cas d'utilisation", 1800, 1150)
    rounded_box(d, (70, 470, 260, 570), "Candidat", PALE_BLUE, BLUE, DARK_BLUE, font(24, True))
    rounded_box(d, (1540, 470, 1730, 570), "Admin FDS", PALE_BLUE, BLUE, DARK_BLUE, font(24, True))
    d.rounded_rectangle((360, 140, 1440, 1010), radius=24, fill="#FAFBFD", outline=f"#{GRID}", width=3)
    label(d, (390, 165), "FDS Portail", 28, DARK_BLUE, True)
    left = [
        ("Consulter les cursus", 250),
        ("Soumettre une candidature", 360),
        ("Simuler le paiement", 470),
        ("Téléverser les documents", 580),
        ("Suivre le dossier", 690),
        ("Visualiser la progression", 800),
        ("Remplacer un document rejeté", 910),
    ]
    right = [
        ("Se connecter", 330),
        ("Consulter les candidatures", 470),
        ("Valider / rejeter", 610),
        ("Auditer les décisions", 750),
    ]
    for text, y in left:
        rounded_box(d, (430, y, 820, y + 72), text, "FFFFFF", BLUE, INK, font(22))
        arrow(d, (260, 520), (430, y + 36))
    for text, y in right:
        rounded_box(d, (980, y, 1370, y + 72), text, "FFFFFF", DARK_BLUE, INK, font(22))
        arrow(d, (1540, 520), (1370, y + 36))
    arrow(d, (820, 396), (980, 506), MUTED)
    label(d, (850, 430), "inclut", 18, MUTED)
    arrow(d, (820, 726), (430, 946), RED)
    label(d, (555, 850), "si rejet", 18, RED)
    arrow(d, (1175, 402), (1175, 470), MUTED)
    label(d, (1200, 430), "protège", 18, MUTED)
    arrow(d, (1175, 682), (1175, 750), MUTED)
    label(d, (1200, 705), "trace", 18, MUTED)
    return save(img, "diagramme_01_cas_utilisation.png")


def diagram_activity() -> Path:
    img, d = canvas("Figure 2 - Diagramme d'activités : candidature et suivi", 1900, 1250)
    items = [
        ("Début", 80, 170, "E8EEF5"),
        ("Consulter les cursus", 310, 170, "FFFFFF"),
        ("Remplir le formulaire", 600, 170, "FFFFFF"),
        ("Paiement simulé", 890, 170, "FFFFFF"),
        ("Upload documents", 1180, 170, "FFFFFF"),
        ("Contrôle fichier", 1470, 170, "FFF7D6"),
        ("Stockage Cloudinary", 1470, 360, "FFFFFF"),
        ("Enregistrement BDD", 1180, 360, "FFFFFF"),
        ("Référence + email", 890, 360, "FFFFFF"),
        ("Suivi avec progression", 600, 360, "E8EEF5"),
        ("Examen admin", 310, 560, "FFFFFF"),
        ("Décision admin", 600, 560, "FFF7D6"),
        ("Document validé", 890, 520, "E7F6ED"),
        ("Document rejeté", 890, 650, "FDE8E6"),
        ("Correction requise", 1180, 650, "FDE8E6"),
        ("Remplacement candidat", 1470, 650, "FFFFFF"),
        ("Dossier validé", 1180, 900, "E7F6ED"),
        ("Fin", 1470, 900, "E8EEF5"),
    ]
    boxes = {}
    for text, x, y, fill in items:
        boxes[text] = (x, y, x + 220, y + 78)
        rounded_box(d, boxes[text], text, fill, BLUE if fill != "FDE8E6" else RED, INK, font(20))
    links = [
        ("Début", "Consulter les cursus"), ("Consulter les cursus", "Remplir le formulaire"),
        ("Remplir le formulaire", "Paiement simulé"), ("Paiement simulé", "Upload documents"),
        ("Upload documents", "Contrôle fichier"), ("Contrôle fichier", "Stockage Cloudinary"),
        ("Stockage Cloudinary", "Enregistrement BDD"), ("Enregistrement BDD", "Référence + email"),
        ("Référence + email", "Suivi avec progression"), ("Suivi avec progression", "Examen admin"),
        ("Examen admin", "Décision admin"), ("Décision admin", "Document validé"),
        ("Décision admin", "Document rejeté"), ("Document rejeté", "Correction requise"),
        ("Correction requise", "Remplacement candidat"), ("Dossier validé", "Fin"),
    ]
    for a, b in links:
        ax1, ay1, ax2, ay2 = boxes[a]
        bx1, by1, bx2, by2 = boxes[b]
        arrow(d, ((ax1 + ax2) // 2 if abs(bx1 - ax1) < 80 else ax2, (ay1 + ay2) // 2), (bx1 if bx1 > ax1 else (bx1 + bx2) // 2, (by1 + by2) // 2))
    arrow(d, (1690, 690), (1580, 248), RED)
    label(d, (1570, 470), "nouvel upload", 18, RED, True)
    arrow(d, (1110, 560), (1180, 936), GREEN)
    label(d, (1080, 740), "si tous validés", 18, GREEN, True)
    label(d, (1510, 275), "Non valide : retour upload", 18, RED)
    return save(img, "diagramme_02_activites.png")


def diagram_sequence() -> Path:
    img, d = canvas("Figure 3 - Diagramme de séquence : soumission, upload et notification", 1900, 1250)
    actors = [("Candidat", 160), ("React", 430), ("FastAPI", 720), ("PostgreSQL", 1010), ("Cloudinary", 1300), ("Resend", 1590)]
    for name, x in actors:
        rounded_box(d, (x - 90, 130, x + 90, 190), name, PALE_BLUE, BLUE, DARK_BLUE, font(20, True))
        d.line((x, 190, x, 1130), fill=f"#{GRID}", width=3)
    messages = [
        (160, 430, 250, "Remplit formulaire"),
        (430, 720, 320, "POST /api/candidature"),
        (720, 1010, 390, "INSERT candidat + référence"),
        (720, 430, 470, "201 + référence"),
        (430, 720, 550, "POST /api/upload"),
        (720, 1300, 630, "Upload sécurisé"),
        (720, 1010, 710, "INSERT/UPDATE document"),
        (720, 1590, 790, "Email confirmation"),
        (430, 720, 890, "GET /api/candidature/{ref}"),
        (720, 1010, 970, "SELECT dossier + statuts"),
        (720, 430, 1050, "Données de suivi + progression"),
    ]
    for x1, x2, y, text in messages:
        arrow(d, (x1, y), (x2, y), DARK_BLUE if x1 < x2 else MUTED)
        label(d, (min(x1, x2) + 15, y - 30), text, 18, INK)
    rounded_box(d, (1450, 1110, 1760, 1180), "Admin : PUT statut\n+ email validation/rejet", "FFFFFF", RED, INK, font(19))
    return save(img, "diagramme_03_sequence.png")


def diagram_classes() -> Path:
    img, d = canvas("Figure 4 - Diagramme de classes UML", 1800, 1150)
    classes = [
        ("Utilisateur", ["id", "email", "mot_de_passe_hash", "role", "created_at"], 120, 200),
        ("Candidat", ["id", "reference_dossier", "nom", "prenom", "email", "statut_paiement", "reference_paiement", "notifications_actives"], 600, 160),
        ("DocumentRequis", ["id", "nom", "description", "format_accepte", "est_obligatoire"], 1150, 200),
        ("DocumentSoumis", ["id", "candidat_id", "document_requis_id", "fichier_url", "statut_validation", "soumis_le", "valide_par", "date_validation"], 620, 690),
    ]
    coords = {}
    for title, attrs, x, y in classes:
        w, h = 400, 70 + len(attrs) * 35
        coords[title] = (x, y, x + w, y + h)
        d.rounded_rectangle((x, y, x + w, y + h), radius=10, fill="#FFFFFF", outline=f"#{BLUE}", width=3)
        d.rectangle((x, y, x + w, y + 54), fill=f"#{PALE_BLUE}", outline=f"#{BLUE}", width=2)
        label(d, (x + 16, y + 13), title, 23, DARK_BLUE, True)
        yy = y + 68
        for attr in attrs:
            label(d, (x + 20, yy), f"+ {attr}", 18, INK)
            yy += 35
    arrow(d, (800, 430), (800, 690), DARK_BLUE)
    label(d, (815, 560), "1 possède N", 18, DARK_BLUE, True)
    arrow(d, (1150, 500), (1020, 760), DARK_BLUE)
    label(d, (1055, 625), "1 définit N", 18, DARK_BLUE, True)
    arrow(d, (520, 390), (650, 780), DARK_BLUE)
    label(d, (455, 590), "1 valide N", 18, DARK_BLUE, True)
    return save(img, "diagramme_04_classes.png")


def diagram_components() -> Path:
    img, d = canvas("Figure 5 - Diagramme de composants", 1900, 1200)
    rounded_box(d, (70, 510, 250, 590), "Candidat", PALE_BLUE, BLUE, DARK_BLUE, font(22, True))
    rounded_box(d, (70, 660, 250, 740), "Admin FDS", PALE_BLUE, BLUE, DARK_BLUE, font(22, True))
    d.rounded_rectangle((360, 140, 1350, 1010), radius=24, fill="#FAFBFD", outline=f"#{GRID}", width=3)
    label(d, (400, 170), "Plateforme FDS - FDS Portail", 28, DARK_BLUE, True)
    rounded_box(d, (470, 270, 820, 360), "Frontend React / Vite", "FFFFFF", BLUE, INK, font(22, True))
    rounded_box(d, (900, 270, 1250, 360), "API FastAPI", "FFFFFF", BLUE, INK, font(22, True))
    modules = [("Auth\nJWT + RBAC", 455, 500), ("Candidature\ncréation + suivi", 705, 500), ("Documents\nupload + remplacement", 955, 500), ("Administration\nvalidation + audit", 705, 680)]
    for text, x, y in modules:
        rounded_box(d, (x, y, x + 230, y + 105), text, PALE_BLUE, DARK_BLUE, INK, font(20))
        arrow(d, (1075, 360), (x + 115, y), MUTED)
    rounded_box(d, (650, 860, 1060, 950), "PostgreSQL\nsource de vérité", "FFFFFF", GREEN, INK, font(22, True))
    for x, y in [(570, 605), (820, 605), (1070, 605), (820, 785)]:
        arrow(d, (x, y), (855, 860), GREEN)
    rounded_box(d, (1470, 360, 1770, 450), "Cloudinary\nstockage fichiers", "FFFFFF", BLUE, INK, font(22))
    rounded_box(d, (1470, 540, 1770, 630), "Resend\nemails", "FFFFFF", BLUE, INK, font(22))
    rounded_box(d, (1470, 760, 1770, 850), "FDS Pay\nphase future", "FFFFFF", GOLD, INK, font(22))
    arrow(d, (250, 550), (470, 315))
    arrow(d, (250, 700), (470, 315))
    arrow(d, (820, 315), (900, 315))
    arrow(d, (1185, 552), (1470, 405))
    arrow(d, (1075, 730), (1470, 585))
    arrow(d, (820, 552), (1470, 805), GOLD)
    return save(img, "diagramme_05_composants.png")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def style_table(table, widths: List[int] | None = None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths and j < len(widths):
                set_cell_width(cell, widths[j])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, LIGHT)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def add_table(doc: Document, rows: List[List[str]]):
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    for i, row in enumerate(rows):
        for j in range(max_cols):
            table.cell(i, j).text = row[j].strip() if j < len(row) else ""
    widths = None
    if max_cols == 2:
        widths = [2600, 6760]
    elif max_cols == 3:
        widths = [2200, 3580, 3580]
    elif max_cols == 4:
        widths = [1700, 2660, 2500, 2500]
    style_table(table, widths)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for line in code.splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_para(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    text = text.strip()
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
        else:
            p.add_run(part.replace("—", "-"))
    return p


def create_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc: Document):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.1))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Cahier des Charges")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FDS Portail - Module 2")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    meta = [
        ("Projet", "Portail public et plateforme de candidature en ligne de la Faculté des Sciences"),
        ("Institution", "Faculté des Sciences - Université d'Etat d'Haïti"),
        ("Version", "Livrable académique"),
        ("Date", "Mai 2026"),
    ]
    doc.add_paragraph()
    table = doc.add_table(rows=len(meta), cols=2)
    for i, (k, v) in enumerate(meta):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
    style_table(table, [2200, 7160])
    doc.add_paragraph()
    callout = doc.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = callout.add_run("Document d'analyse, de conception et de cadrage fonctionnel du MVP.")
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_toc(doc: Document):
    doc.add_heading("Table des matières", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Dans Microsoft Word : Références > Mettre à jour la table pour générer automatiquement les numéros de page.")
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    for title in [
        "1. Problème Observé", "2. Solution Proposée", "3. Argumentation",
        "4. Priorisation MoSCoW", "5. MVP et Walking Skeleton",
        "6. Use Cases et User Stories", "7. Scénarios et Séquences",
        "8. Modèle de Données", "9. Architecture et Composants",
        "10. Choix Technologiques", "11. Validation, Risques et Limites",
    ]:
        doc.add_paragraph(title, style="List Bullet")
    doc.add_page_break()


def markdown_to_docx(doc: Document, md: str, figures: List[Path]):
    lines = md.splitlines()
    in_code = False
    code_lang = ""
    code_lines: List[str] = []
    figure_index = 0
    table_rows: List[List[str]] = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            clean = [row for row in table_rows if not all(re.fullmatch(r"-+", c.strip()) for c in row)]
            add_table(doc, clean)
            table_rows = []

    for line in lines:
        raw = line.rstrip()
        if raw.startswith("```"):
            if not in_code:
                flush_table()
                in_code = True
                code_lang = raw[3:].strip()
                code_lines = []
            else:
                if code_lang == "mermaid" and figure_index < len(figures):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(figures[figure_index]), width=Inches(6.5))
                    figure_index += 1
                else:
                    add_code_block(doc, "\n".join(code_lines))
                in_code = False
            continue
        if in_code:
            code_lines.append(raw)
            continue

        if not raw.strip():
            flush_table()
            continue
        if raw.strip() == "---":
            flush_table()
            continue
        if raw.startswith("# "):
            flush_table()
            continue
        if raw.startswith("## "):
            flush_table()
            text = raw[3:].replace("§", "").strip()
            doc.add_heading(text, level=1)
            continue
        if raw.startswith("### "):
            flush_table()
            text = raw[4:].strip()
            doc.add_heading(text, level=2)
            continue
        if raw.startswith("#### "):
            flush_table()
            text = raw[5:].strip()
            doc.add_heading(text, level=3)
            continue
        if raw.lstrip().startswith("- "):
            flush_table()
            add_para(doc, raw.lstrip()[2:], style="List Bullet")
            continue
        if re.match(r"^\d+\. ", raw.strip()):
            flush_table()
            add_para(doc, re.sub(r"^\d+\. ", "", raw.strip()), style="List Number")
            continue
        if raw.startswith(">"):
            flush_table()
            p = add_para(doc, raw.lstrip("> ").strip())
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor.from_string(MUTED)
            continue
        if raw.startswith("|") and raw.endswith("|"):
            cells = [cell.strip() for cell in raw.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                continue
            table_rows.append(cells)
            continue
        flush_table()
        add_para(doc, raw)
    flush_table()


def add_footer(doc: Document):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("FDS Portail - Cahier des charges - Module 2")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=26, leading=32, alignment=TA_CENTER, textColor=colors.HexColor(f"#{DARK_BLUE}"),
        spaceAfter=14,
    ))
    styles.add(ParagraphStyle(
        name="CoverSub", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=16, leading=20, alignment=TA_CENTER, textColor=colors.HexColor(f"#{BLUE}"),
        spaceAfter=20,
    ))
    styles.add(ParagraphStyle(
        name="H1FDS", parent=styles["Heading1"], fontName="Helvetica-Bold",
        fontSize=16, leading=20, textColor=colors.HexColor(f"#{BLUE}"),
        spaceBefore=14, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="H2FDS", parent=styles["Heading2"], fontName="Helvetica-Bold",
        fontSize=13, leading=16, textColor=colors.HexColor(f"#{BLUE}"),
        spaceBefore=10, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="H3FDS", parent=styles["Heading3"], fontName="Helvetica-Bold",
        fontSize=11.5, leading=14, textColor=colors.HexColor(f"#{DARK_BLUE}"),
        spaceBefore=8, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="BodyFDS", parent=styles["BodyText"], fontName="Helvetica",
        fontSize=9.5, leading=12.3, textColor=colors.HexColor(f"#{INK}"),
        spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="Muted", parent=styles["BodyText"], fontName="Helvetica-Oblique",
        fontSize=9.2, leading=12, textColor=colors.HexColor(f"#{MUTED}"),
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="CodeFDS", parent=styles["Code"], fontName="Courier",
        fontSize=7.6, leading=9.2, leftIndent=10, rightIndent=10,
        backColor=colors.HexColor("#F6F8FA"), spaceBefore=4, spaceAfter=8,
    ))
    return styles


def html_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("**", "")
        .replace("`", "")
    )


def pdf_table(rows: List[List[str]]) -> Table:
    max_cols = max(len(row) for row in rows)
    data = []
    style = pdf_styles()
    for row in rows:
        data.append([Paragraph(html_escape(row[i].strip()) if i < len(row) else "", style["BodyFDS"]) for i in range(max_cols)])
    widths_map = {
        2: [1.8 * inch, 4.5 * inch],
        3: [1.45 * inch, 2.45 * inch, 2.4 * inch],
        4: [1.2 * inch, 1.75 * inch, 1.65 * inch, 1.7 * inch],
    }
    table = Table(data, colWidths=widths_map.get(max_cols), hAlign="CENTER", repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{GRID}")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{DARK_BLUE}")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def markdown_to_pdf_story(md: str, figures: List[Path]):
    styles = pdf_styles()
    story = []
    if LOGO.exists():
        story.append(RLImage(str(LOGO), width=1.0 * inch, height=1.0 * inch, hAlign="CENTER"))
        story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Cahier des Charges", styles["CoverTitle"]))
    story.append(Paragraph("FDS Portail - Module 2", styles["CoverSub"]))
    story.append(pdf_table([
        ["Projet", "Portail public et plateforme de candidature en ligne de la Faculté des Sciences"],
        ["Institution", "Faculté des Sciences - Université d'Etat d'Haïti"],
        ["Version", "Livrable académique"],
        ["Date", "Mai 2026"],
    ]))
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph("Document d'analyse, de conception et de cadrage fonctionnel du MVP.", styles["Muted"]))
    story.append(PageBreak())
    story.append(Paragraph("Table des matières", styles["H1FDS"]))
    for title in [
        "1. Problème Observé", "2. Solution Proposée", "3. Argumentation",
        "4. Priorisation MoSCoW", "5. MVP et Walking Skeleton",
        "6. Use Cases et User Stories", "7. Scénarios et Séquences",
        "8. Modèle de Données", "9. Architecture et Composants",
        "10. Choix Technologiques", "11. Validation, Risques et Limites",
    ]:
        story.append(Paragraph(f"• {title}", styles["BodyFDS"]))
    story.append(PageBreak())

    lines = md.splitlines()
    in_code = False
    code_lang = ""
    code_lines: List[str] = []
    figure_index = 0
    table_rows: List[List[str]] = []
    bullets: List[ListItem] = []
    numbers: List[ListItem] = []

    def flush_lists():
        nonlocal bullets, numbers
        if bullets:
            story.append(ListFlowable(bullets, bulletType="bullet", leftIndent=18))
            bullets = []
        if numbers:
            story.append(ListFlowable(numbers, bulletType="1", leftIndent=18))
            numbers = []

    def flush_table():
        nonlocal table_rows
        flush_lists()
        if table_rows:
            story.append(pdf_table(table_rows))
            story.append(Spacer(1, 6))
            table_rows = []

    for line in lines:
        raw = line.rstrip()
        if raw.startswith("```"):
            if not in_code:
                flush_table()
                in_code = True
                code_lang = raw[3:].strip()
                code_lines = []
            else:
                if code_lang == "mermaid" and figure_index < len(figures):
                    story.append(RLImage(str(figures[figure_index]), width=6.4 * inch, height=4.1 * inch, hAlign="CENTER"))
                    story.append(Spacer(1, 8))
                    figure_index += 1
                else:
                    story.append(Paragraph("<br/>".join(html_escape(x) for x in code_lines), styles["CodeFDS"]))
                in_code = False
            continue
        if in_code:
            code_lines.append(raw)
            continue
        if not raw.strip() or raw.strip() == "---":
            flush_table()
            continue
        if raw.startswith("# "):
            flush_table()
            continue
        if raw.startswith("## "):
            flush_table()
            story.append(Paragraph(html_escape(raw[3:].replace("§", "").strip()), styles["H1FDS"]))
            continue
        if raw.startswith("### "):
            flush_table()
            story.append(Paragraph(html_escape(raw[4:].strip()), styles["H2FDS"]))
            continue
        if raw.startswith("#### "):
            flush_table()
            story.append(Paragraph(html_escape(raw[5:].strip()), styles["H3FDS"]))
            continue
        if raw.lstrip().startswith("- "):
            item = Paragraph(html_escape(raw.lstrip()[2:]), styles["BodyFDS"])
            bullets.append(ListItem(item))
            continue
        if re.match(r"^\d+\. ", raw.strip()):
            item = Paragraph(html_escape(re.sub(r"^\d+\. ", "", raw.strip())), styles["BodyFDS"])
            numbers.append(ListItem(item))
            continue
        if raw.startswith(">"):
            flush_table()
            story.append(Paragraph(html_escape(raw.lstrip("> ").strip()), styles["Muted"]))
            continue
        if raw.startswith("|") and raw.endswith("|"):
            cells = [cell.strip() for cell in raw.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                continue
            table_rows.append(cells)
            continue
        flush_table()
        story.append(Paragraph(html_escape(raw), styles["BodyFDS"]))
    flush_table()
    return story


def add_pdf_footer(canvas_obj, doc):
    canvas_obj.saveState()
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(colors.HexColor(f"#{MUTED}"))
    canvas_obj.drawCentredString(letter[0] / 2, 0.45 * inch, f"FDS Portail - Cahier des charges - Page {doc.page}")
    canvas_obj.restoreState()


def build_pdf(md: str, figures: List[Path]):
    pdf = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=1 * inch,
        leftMargin=1 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.8 * inch,
    )
    pdf.build(markdown_to_pdf_story(md, figures), onFirstPage=add_pdf_footer, onLaterPages=add_pdf_footer)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    figures = [
        diagram_use_cases(),
        diagram_activity(),
        diagram_sequence(),
        diagram_classes(),
        diagram_components(),
    ]
    doc = Document()
    create_styles(doc)
    add_cover(doc)
    add_toc(doc)
    md = SOURCE.read_text(encoding="utf-8")
    markdown_to_docx(doc, md, figures)
    add_footer(doc)
    doc.save(DOCX_OUT)
    build_pdf(md, figures)
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()
